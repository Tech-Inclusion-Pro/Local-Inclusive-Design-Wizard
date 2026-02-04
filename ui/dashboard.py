"""Dashboard screen with session management."""

import asyncio
from datetime import datetime
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QFrame, QScrollArea, QGridLayout, QDialog, QLineEdit,
    QComboBox, QMessageBox, QSizePolicy, QTextEdit, QProgressBar,
    QFileDialog
)
from PyQt6.QtCore import Qt, pyqtSignal, QThread, pyqtSlot
from PyQt6.QtGui import QFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import COLORS
from prompts.system_prompts import CONSULTATION_TYPES
from core.database import DatabaseManager, Session, ConsultOverview
from core.auth import AuthManager


class OverviewWorker(QThread):
    """Worker thread for generating AI overview."""

    chunk_received = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, ai_manager, conversations_text: str):
        super().__init__()
        self.ai_manager = ai_manager
        self.conversations_text = conversations_text
        self.full_response = ""

    def run(self):
        """Run the AI analysis."""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        try:
            prompt = f"""Analyze the following consultation conversations and provide a comprehensive overview report.

CONVERSATIONS DATA:
{self.conversations_text}

Please provide a structured report with the following sections:

## 1. Common Support Questions
Identify and list the most common questions the educator is asking for support. What patterns do you see in what they need help with?

## 2. Student Support Themes
What common statements or concerns does the educator express about supporting their students? What are their main priorities regarding student needs?

## 3. Universal Design Questions
What questions is the educator asking about making their materials more universally designed? What UDL principles are they most curious about?

## 4. Accessibility & Technology Questions
What questions does the educator have about accessibility and technology? What technical accessibility concerns come up most often?

## 5. Learning Progress Assessment
Based on these conversations, assess the educator's understanding of:
- UDL (Universal Design for Learning) principles
- WCAG (Web Content Accessibility Guidelines) standards
Rate their progress and identify areas of strength and areas needing more development.

## 6. Personalized Recommendations
Based on the analysis above, provide 5-7 specific, actionable recommendations for this educator to improve their inclusive design practice.

## 7. Suggested Next Steps
What should this educator focus on learning next? Provide specific resources or topics to explore.

Format the response in clear markdown with headers and bullet points for readability."""

            async def generate():
                try:
                    async for chunk in self.ai_manager.generate_response(prompt, None):
                        self.full_response += chunk
                        self.chunk_received.emit(chunk)
                except Exception as e:
                    self.error.emit(str(e))
                    return

            loop.run_until_complete(generate())
            self.finished.emit(self.full_response)

        except Exception as e:
            self.error.emit(str(e))
        finally:
            loop.close()


class ConsultOverviewDialog(QDialog):
    """Dialog for viewing AI-generated overview of all consultations."""

    def __init__(self, db_manager: DatabaseManager, auth_manager: AuthManager,
                 ai_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.auth = auth_manager
        self.ai = ai_manager
        self.overview_content = ""
        self.worker = None
        self.generated_time = None
        self.sessions_count = 0
        self.saved_overview_id = None

        self.setWindowTitle("Consultation Overview")
        self.setModal(True)
        self.setMinimumSize(700, 600)
        self.setup_ui()

    def setup_ui(self):
        """Set up the overview dialog UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
            QTextEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
            }}
            QProgressBar {{
                border: none;
                border-radius: 4px;
                background-color: {COLORS['dark_input']};
                text-align: center;
            }}
            QProgressBar::chunk {{
                background-color: {COLORS['primary']};
                border-radius: 4px;
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Header
        header_layout = QHBoxLayout()

        title = QLabel("Consultation Overview")
        title.setFont(QFont("Arial", 22, QFont.Weight.Bold))
        header_layout.addWidget(title)

        header_layout.addStretch()

        # Timestamp
        self.timestamp_label = QLabel("")
        self.timestamp_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 12px;")
        header_layout.addWidget(self.timestamp_label)

        layout.addLayout(header_layout)

        # Description
        desc = QLabel(
            "This overview analyzes all your consultations to identify patterns, "
            "assess your learning progress, and provide personalized recommendations."
        )
        desc.setWordWrap(True)
        desc.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 13px;")
        layout.addWidget(desc)

        # Progress bar (hidden initially)
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)  # Indeterminate
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.hide()
        layout.addWidget(self.progress_bar)

        # Status label
        self.status_label = QLabel("")
        self.status_label.setStyleSheet(f"color: {COLORS['primary']}; font-size: 13px;")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.status_label)

        # Content area
        self.content_area = QTextEdit()
        self.content_area.setReadOnly(True)
        self.content_area.setPlaceholderText(
            "Click 'Generate Overview' to analyze your consultations..."
        )
        self.content_area.setMinimumHeight(350)
        layout.addWidget(self.content_area)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        # Generate button
        self.generate_btn = QPushButton("Generate Overview")
        self.generate_btn.clicked.connect(self.generate_overview)
        self.generate_btn.setMinimumHeight(44)
        self.generate_btn.setMinimumWidth(150)
        btn_layout.addWidget(self.generate_btn)

        btn_layout.addStretch()

        # Export button
        self.export_btn = QPushButton("Export Document")
        self.export_btn.setProperty("class", "secondary")
        self.export_btn.clicked.connect(self.export_overview)
        self.export_btn.setMinimumHeight(44)
        self.export_btn.setEnabled(False)
        btn_layout.addWidget(self.export_btn)

        # Close button
        close_btn = QPushButton("Close")
        close_btn.setProperty("class", "secondary")
        close_btn.clicked.connect(self.accept)
        close_btn.setMinimumHeight(44)
        btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)

    def generate_overview(self):
        """Generate the AI overview of all consultations."""
        user = self.auth.get_current_user()
        if not user:
            QMessageBox.warning(self, "Error", "No user logged in.")
            return

        # Get all sessions for the user
        sessions = self.db.get_user_sessions(user.id)

        if not sessions:
            QMessageBox.information(
                self,
                "No Consultations",
                "You don't have any consultations yet. Start a consultation first!"
            )
            return

        # Store sessions count for saving
        self.sessions_count = len(sessions)

        # Gather all conversation data
        conversations_text = self._gather_conversations(sessions)

        if not conversations_text.strip():
            QMessageBox.information(
                self,
                "No Conversation Data",
                "Your consultations don't have any conversation data yet."
            )
            return

        # Update UI for generating state
        self.generate_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        self.progress_bar.show()
        self.status_label.setText("Analyzing your consultations...")
        self.content_area.clear()
        self.overview_content = ""

        # Start worker thread
        self.worker = OverviewWorker(self.ai, conversations_text)
        self.worker.chunk_received.connect(self.on_chunk_received)
        self.worker.finished.connect(self.on_generation_finished)
        self.worker.error.connect(self.on_generation_error)
        self.worker.start()

    def _gather_conversations(self, sessions) -> str:
        """Gather all conversation text from sessions."""
        all_conversations = []

        for session in sessions:
            # Reload session to get conversation data
            full_session = self.db.get_session_by_id(session.id)
            if not full_session:
                continue

            conversation = full_session.conversation
            if not conversation:
                continue

            session_text = f"\n--- Consultation: {full_session.title} ({full_session.template_type}) ---\n"
            session_text += f"Created: {full_session.created_at.strftime('%Y-%m-%d')}\n\n"

            for msg in conversation:
                role = "Educator" if msg["role"] == "user" else "AI Consultant"
                content = msg.get("content", "")
                session_text += f"{role}: {content}\n\n"

            all_conversations.append(session_text)

        return "\n".join(all_conversations)

    @pyqtSlot(str)
    def on_chunk_received(self, chunk: str):
        """Handle received chunk from AI."""
        self.overview_content += chunk
        self.content_area.setPlainText(self.overview_content)
        # Scroll to bottom
        scrollbar = self.content_area.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    @pyqtSlot(str)
    def on_generation_finished(self, full_response: str):
        """Handle generation completion."""
        self.progress_bar.hide()
        self.generate_btn.setEnabled(True)
        self.export_btn.setEnabled(True)
        self.status_label.setText("Overview generated and saved!")
        self.status_label.setStyleSheet(f"color: {COLORS['success']}; font-size: 13px;")

        # Set timestamp
        self.generated_time = datetime.now()
        self.timestamp_label.setText(
            f"Generated: {self.generated_time.strftime('%B %d, %Y at %I:%M %p')}"
        )

        # Save to database
        user = self.auth.get_current_user()
        if user:
            title = f"Overview - {self.generated_time.strftime('%B %d, %Y at %I:%M %p')}"
            overview = self.db.create_overview(
                user_id=user.id,
                title=title,
                content=self.overview_content,
                sessions_analyzed=self.sessions_count
            )
            self.saved_overview_id = overview.id

    @pyqtSlot(str)
    def on_generation_error(self, error_msg: str):
        """Handle generation error."""
        self.progress_bar.hide()
        self.generate_btn.setEnabled(True)
        self.status_label.setText(f"Error: {error_msg}")
        self.status_label.setStyleSheet(f"color: {COLORS['error']}; font-size: 13px;")

        QMessageBox.warning(
            self,
            "Generation Error",
            f"Failed to generate overview: {error_msg}\n\n"
            "Please check your AI settings and try again."
        )

    def export_overview(self):
        """Export the overview to an accessible Word document."""
        if not self.overview_content:
            QMessageBox.warning(self, "No Content", "Generate an overview first.")
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Export Overview",
            f"consultation_overview_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Document (*.docx)"
        )

        if not filepath:
            return

        if not filepath.endswith(".docx"):
            filepath += ".docx"

        try:
            self._create_accessible_document(filepath)
            QMessageBox.information(
                self,
                "Export Successful",
                f"Overview exported to:\n{filepath}"
            )
        except Exception as e:
            QMessageBox.warning(
                self,
                "Export Failed",
                f"Failed to export: {str(e)}"
            )

    def _create_accessible_document(self, filepath: str):
        """Create a WCAG-accessible Word document."""
        doc = Document()

        # Set document properties for accessibility
        doc.core_properties.title = "Consultation Overview - Inclusive Design Wizard"
        doc.core_properties.subject = "Learning Progress and Recommendations"
        doc.core_properties.author = "Inclusive Design Wizard"

        # Title
        title = doc.add_heading("Consultation Overview", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        if self.generated_time:
            timestamp_para = doc.add_paragraph()
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = timestamp_para.add_run(
                f"Generated: {self.generated_time.strftime('%B %d, %Y at %I:%M %p')}"
            )
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # Introduction
        intro = doc.add_paragraph()
        intro.add_run(
            "This document provides an AI-generated analysis of your inclusive design "
            "consultations, including patterns in your questions, learning progress assessment, "
            "and personalized recommendations for improvement."
        )

        doc.add_paragraph()

        # Parse and add content with proper headings
        self._parse_markdown_to_docx(doc, self.overview_content)

        # Add accessibility statement
        doc.add_page_break()
        doc.add_heading("Accessibility Statement", level=1)
        accessibility_text = doc.add_paragraph()
        accessibility_text.add_run(
            "This document was created with accessibility in mind following WCAG guidelines:\n\n"
            "• Proper heading hierarchy for screen reader navigation\n"
            "• Sufficient color contrast for text\n"
            "• Clear, descriptive section titles\n"
            "• Logical reading order\n"
            "• No images requiring alternative text\n\n"
            "If you need this document in an alternative format, please contact your "
            "accessibility support team."
        )

        # Footer with generation info
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f"Generated by Inclusive Design Wizard | {datetime.now().strftime('%Y-%m-%d')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.save(filepath)

    def _parse_markdown_to_docx(self, doc, markdown_text: str):
        """Parse markdown-style text and add to document with proper formatting."""
        lines = markdown_text.split('\n')
        current_para = None
        in_list = False

        for line in lines:
            line = line.strip()

            if not line:
                current_para = None
                in_list = False
                continue

            # Handle headings
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
                current_para = None
                in_list = False
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
                current_para = None
                in_list = False
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                current_para = None
                in_list = False
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
                in_list = True
                current_para = None
            elif line.startswith('• '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
                in_list = True
                current_para = None
            # Handle numbered lists
            elif len(line) > 2 and line[0].isdigit() and line[1] in '.):':
                para = doc.add_paragraph(line[2:].strip(), style='List Number')
                in_list = True
                current_para = None
            # Handle bold text (simple)
            elif line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
                current_para = None
            # Regular paragraph
            else:
                if current_para is None:
                    current_para = doc.add_paragraph()
                    current_para.add_run(line)
                else:
                    current_para.add_run(' ' + line)


class NotesDialog(QDialog):
    """Dialog for viewing and adding notes to a session."""

    note_added = pyqtSignal(int, str)  # session_id, note_content

    def __init__(self, session_id: int, session_title: str, notes: list, parent=None):
        super().__init__(parent)
        self.session_id = session_id
        self.notes = notes
        self.setWindowTitle(f"Notes - {session_title}")
        self.setModal(True)
        self.setMinimumSize(500, 450)
        self.setup_ui()

    def setup_ui(self):
        """Set up the notes dialog UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
            QTextEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                padding: 8px;
                font-size: 14px;
            }}
            QTextEdit:focus {{
                border: 1px solid {COLORS['primary']};
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Title
        title = QLabel("Session Notes")
        title.setFont(QFont("Arial", 20, QFont.Weight.Bold))
        layout.addWidget(title)

        # Existing notes section
        notes_label = QLabel("Previous Notes")
        notes_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(notes_label)

        # Notes scroll area
        notes_scroll = QScrollArea()
        notes_scroll.setWidgetResizable(True)
        notes_scroll.setMaximumHeight(200)
        notes_scroll.setStyleSheet(f"""
            QScrollArea {{
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                background-color: {COLORS['dark_bg']};
            }}
        """)

        notes_container = QWidget()
        notes_layout = QVBoxLayout(notes_container)
        notes_layout.setSpacing(8)
        notes_layout.setContentsMargins(8, 8, 8, 8)

        if self.notes:
            for note in reversed(self.notes):  # Show newest first
                note_frame = QFrame()
                note_frame.setStyleSheet(f"""
                    QFrame {{
                        background-color: {COLORS['dark_input']};
                        border-radius: 6px;
                        padding: 8px;
                    }}
                """)
                note_layout = QVBoxLayout(note_frame)
                note_layout.setSpacing(4)
                note_layout.setContentsMargins(8, 8, 8, 8)

                # Timestamp
                timestamp = datetime.fromisoformat(note['timestamp'])
                time_label = QLabel(timestamp.strftime("%b %d, %Y at %I:%M %p"))
                time_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
                note_layout.addWidget(time_label)

                # Content
                content_label = QLabel(note['content'])
                content_label.setWordWrap(True)
                content_label.setStyleSheet(f"color: {COLORS['text']}; font-size: 13px;")
                note_layout.addWidget(content_label)

                notes_layout.addWidget(note_frame)
        else:
            empty_label = QLabel("No notes yet. Add your first note below!")
            empty_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-style: italic;")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            notes_layout.addWidget(empty_label)

        notes_layout.addStretch()
        notes_scroll.setWidget(notes_container)
        layout.addWidget(notes_scroll)

        # Add new note section
        new_note_label = QLabel("Add New Note")
        new_note_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(new_note_label)

        self.note_input = QTextEdit()
        self.note_input.setPlaceholderText("Type your note here...")
        self.note_input.setAccessibleName("New note input")
        self.note_input.setMinimumHeight(80)
        self.note_input.setMaximumHeight(120)
        layout.addWidget(self.note_input)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        close_btn = QPushButton("Close")
        close_btn.setProperty("class", "secondary")
        close_btn.clicked.connect(self.reject)
        close_btn.setMinimumHeight(44)
        btn_layout.addWidget(close_btn)

        btn_layout.addStretch()

        save_btn = QPushButton("Save Note")
        save_btn.clicked.connect(self.save_note)
        save_btn.setMinimumHeight(44)
        save_btn.setMinimumWidth(120)
        btn_layout.addWidget(save_btn)

        layout.addLayout(btn_layout)

    def save_note(self):
        """Save the new note."""
        content = self.note_input.toPlainText().strip()
        if not content:
            QMessageBox.warning(self, "Empty Note", "Please enter some text for your note.")
            return

        self.note_added.emit(self.session_id, content)
        self.accept()


class OverviewNotesDialog(QDialog):
    """Dialog for viewing and adding notes to an overview."""

    note_added = pyqtSignal(int, str)  # overview_id, note_content

    def __init__(self, overview_id: int, overview_title: str, notes: list, parent=None):
        super().__init__(parent)
        self.overview_id = overview_id
        self.notes = notes
        self.setWindowTitle(f"Notes - {overview_title}")
        self.setModal(True)
        self.setMinimumSize(500, 450)
        self.setup_ui()

    def setup_ui(self):
        """Set up the notes dialog UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
            QTextEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                padding: 8px;
                font-size: 14px;
            }}
            QTextEdit:focus {{
                border: 1px solid {COLORS['primary']};
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Title
        title = QLabel("Overview Notes")
        title.setFont(QFont("Arial", 20, QFont.Weight.Bold))
        layout.addWidget(title)

        # Existing notes section
        notes_label = QLabel("Previous Notes")
        notes_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(notes_label)

        # Notes scroll area
        notes_scroll = QScrollArea()
        notes_scroll.setWidgetResizable(True)
        notes_scroll.setMaximumHeight(200)
        notes_scroll.setStyleSheet(f"""
            QScrollArea {{
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                background-color: {COLORS['dark_bg']};
            }}
        """)

        notes_container = QWidget()
        notes_layout = QVBoxLayout(notes_container)
        notes_layout.setSpacing(8)
        notes_layout.setContentsMargins(8, 8, 8, 8)

        if self.notes:
            for note in reversed(self.notes):
                note_frame = QFrame()
                note_frame.setStyleSheet(f"""
                    QFrame {{
                        background-color: {COLORS['dark_input']};
                        border-radius: 6px;
                        padding: 8px;
                    }}
                """)
                note_layout_inner = QVBoxLayout(note_frame)
                note_layout_inner.setSpacing(4)
                note_layout_inner.setContentsMargins(8, 8, 8, 8)

                timestamp = datetime.fromisoformat(note['timestamp'])
                time_label = QLabel(timestamp.strftime("%b %d, %Y at %I:%M %p"))
                time_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
                note_layout_inner.addWidget(time_label)

                content_label = QLabel(note['content'])
                content_label.setWordWrap(True)
                content_label.setStyleSheet(f"color: {COLORS['text']}; font-size: 13px;")
                note_layout_inner.addWidget(content_label)

                notes_layout.addWidget(note_frame)
        else:
            empty_label = QLabel("No notes yet. Add your first note below!")
            empty_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-style: italic;")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            notes_layout.addWidget(empty_label)

        notes_layout.addStretch()
        notes_scroll.setWidget(notes_container)
        layout.addWidget(notes_scroll)

        # Add new note section
        new_note_label = QLabel("Add New Note")
        new_note_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(new_note_label)

        self.note_input = QTextEdit()
        self.note_input.setPlaceholderText("Type your note here...")
        self.note_input.setAccessibleName("New note input")
        self.note_input.setMinimumHeight(80)
        self.note_input.setMaximumHeight(120)
        layout.addWidget(self.note_input)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        close_btn = QPushButton("Close")
        close_btn.setProperty("class", "secondary")
        close_btn.clicked.connect(self.reject)
        close_btn.setMinimumHeight(44)
        btn_layout.addWidget(close_btn)

        btn_layout.addStretch()

        save_btn = QPushButton("Save Note")
        save_btn.clicked.connect(self.save_note)
        save_btn.setMinimumHeight(44)
        save_btn.setMinimumWidth(120)
        btn_layout.addWidget(save_btn)

        layout.addLayout(btn_layout)

    def save_note(self):
        """Save the new note."""
        content = self.note_input.toPlainText().strip()
        if not content:
            QMessageBox.warning(self, "Empty Note", "Please enter some text for your note.")
            return

        self.note_added.emit(self.overview_id, content)
        self.accept()


class ViewOverviewDialog(QDialog):
    """Dialog for viewing a saved overview."""

    def __init__(self, overview: ConsultOverview, parent=None):
        super().__init__(parent)
        self.overview = overview
        self.setWindowTitle(overview.title)
        self.setModal(True)
        self.setMinimumSize(700, 550)
        self.setup_ui()

    def setup_ui(self):
        """Set up the view dialog UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
            QTextEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Header
        header_layout = QHBoxLayout()

        title = QLabel(self.overview.title)
        title.setFont(QFont("Arial", 20, QFont.Weight.Bold))
        header_layout.addWidget(title)

        header_layout.addStretch()

        # Metadata
        meta_layout = QVBoxLayout()
        created_label = QLabel(f"Created: {self.overview.created_at.strftime('%B %d, %Y at %I:%M %p')}")
        created_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
        meta_layout.addWidget(created_label)

        sessions_label = QLabel(f"Sessions analyzed: {self.overview.sessions_analyzed}")
        sessions_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
        meta_layout.addWidget(sessions_label)

        header_layout.addLayout(meta_layout)
        layout.addLayout(header_layout)

        # Content
        content_area = QTextEdit()
        content_area.setReadOnly(True)
        content_area.setPlainText(self.overview.content)
        layout.addWidget(content_area)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        export_btn = QPushButton("Export Document")
        export_btn.setProperty("class", "secondary")
        export_btn.clicked.connect(self.export_overview)
        export_btn.setMinimumHeight(44)
        btn_layout.addWidget(export_btn)

        btn_layout.addStretch()

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        close_btn.setMinimumHeight(44)
        close_btn.setMinimumWidth(100)
        btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)

    def export_overview(self):
        """Export the overview to a Word document."""
        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Export Overview",
            f"overview_{self.overview.created_at.strftime('%Y%m%d_%H%M')}.docx",
            "Word Document (*.docx)"
        )

        if not filepath:
            return

        if not filepath.endswith(".docx"):
            filepath += ".docx"

        try:
            self._create_document(filepath)
            QMessageBox.information(self, "Export Successful", f"Overview exported to:\n{filepath}")
        except Exception as e:
            QMessageBox.warning(self, "Export Failed", f"Failed to export: {str(e)}")

    def _create_document(self, filepath: str):
        """Create a WCAG-accessible Word document."""
        doc = Document()

        doc.core_properties.title = self.overview.title
        doc.core_properties.subject = "Consultation Overview"

        # Title
        title = doc.add_heading(self.overview.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated: {self.overview.created_at.strftime('%B %d, %Y at %I:%M %p')}\n")
        meta.add_run(f"Sessions Analyzed: {self.overview.sessions_analyzed}")

        doc.add_paragraph()

        # Content
        lines = self.overview.content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(filepath)


class OverviewCard(QFrame):
    """Card widget for displaying a saved overview."""

    open_requested = pyqtSignal(int)
    delete_requested = pyqtSignal(int)
    notes_requested = pyqtSignal(int, str)
    title_changed = pyqtSignal(int, str)

    def __init__(self, overview: ConsultOverview):
        super().__init__()
        self.overview_id = overview.id
        self.overview_title = overview.title
        self.original_title = overview.title
        self.setup_ui(overview)
        self.setFocusPolicy(Qt.FocusPolicy.TabFocus)

    def setup_ui(self, overview: ConsultOverview):
        """Set up the card UI."""
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {COLORS['dark_card']};
                border-radius: 12px;
                padding: 16px;
                border-left: 4px solid {COLORS['secondary']};
            }}
            QFrame:hover {{
                background-color: {COLORS['dark_input']};
            }}
        """)
        self.setMinimumWidth(250)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        layout = QVBoxLayout(self)
        layout.setSpacing(8)
        layout.setContentsMargins(0, 0, 0, 0)

        # Content area
        content_widget = QWidget()
        content_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(6)
        content_layout.setContentsMargins(0, 0, 0, 0)

        # Editable Title
        self.title_input = QLineEdit(overview.title)
        self.title_input.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        self.title_input.setStyleSheet(f"""
            QLineEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 6px;
                padding: 6px 8px;
            }}
            QLineEdit:focus {{
                border: 1px solid {COLORS['primary']};
            }}
        """)
        self.title_input.setAccessibleName("Edit overview title")
        self.title_input.editingFinished.connect(self.on_title_changed)
        content_layout.addWidget(self.title_input)

        # Type badge
        type_label = QLabel("Overview Report")
        type_label.setStyleSheet(f"""
            background-color: {COLORS['secondary']};
            color: {COLORS['text']};
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 11px;
        """)
        type_label.setMaximumWidth(120)
        type_label.setSizePolicy(QSizePolicy.Policy.Maximum, QSizePolicy.Policy.Fixed)
        content_layout.addWidget(type_label)

        # Sessions analyzed
        sessions_label = QLabel(f"Sessions analyzed: {overview.sessions_analyzed}")
        sessions_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 12px;")
        content_layout.addWidget(sessions_label)

        # Created date
        created_label = QLabel(f"Created: {overview.created_at.strftime('%b %d, %Y at %I:%M %p')}")
        created_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
        content_layout.addWidget(created_label)

        layout.addWidget(content_widget)
        layout.addStretch(1)

        # Actions row
        actions_widget = QWidget()
        actions_widget.setFixedHeight(28)
        actions_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        actions = QHBoxLayout(actions_widget)
        actions.setSpacing(6)
        actions.setContentsMargins(0, 4, 0, 0)

        # Open button
        open_btn = QPushButton("Open")
        open_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {COLORS['success']};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 2px 8px;
                font-size: 11px;
                font-weight: bold;
                min-height: 20px;
                max-height: 24px;
            }}
            QPushButton:hover {{
                background-color: #27ae60;
            }}
        """)
        open_btn.setFixedSize(50, 24)
        open_btn.clicked.connect(lambda: self.open_requested.emit(self.overview_id))
        open_btn.setAccessibleName(f"Open {overview.title}")
        actions.addWidget(open_btn)

        # Notes button
        notes_btn = QPushButton("Notes")
        notes_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {COLORS['primary']};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 2px 8px;
                font-size: 11px;
                font-weight: bold;
                min-height: 20px;
                max-height: 24px;
            }}
            QPushButton:hover {{
                background-color: #5849c4;
            }}
        """)
        notes_btn.setFixedSize(55, 24)
        notes_btn.clicked.connect(lambda: self.notes_requested.emit(self.overview_id, self.overview_title))
        notes_btn.setAccessibleName(f"Notes for {overview.title}")
        actions.addWidget(notes_btn)

        actions.addStretch()

        # Delete button
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {COLORS['error']};
                border: 1px solid {COLORS['error']};
                border-radius: 4px;
                padding: 2px 6px;
                font-size: 10px;
                min-height: 18px;
                max-height: 22px;
            }}
            QPushButton:hover {{
                background-color: {COLORS['error']};
                color: white;
            }}
        """)
        delete_btn.setFixedSize(50, 22)
        delete_btn.clicked.connect(lambda: self.delete_requested.emit(self.overview_id))
        delete_btn.setAccessibleName(f"Delete {overview.title}")
        actions.addWidget(delete_btn)

        layout.addWidget(actions_widget)

        self.setAccessibleName(f"{overview.title}. Overview Report. {overview.sessions_analyzed} sessions analyzed.")

    def on_title_changed(self):
        """Handle title edit completion."""
        new_title = self.title_input.text().strip()
        if new_title and new_title != self.original_title:
            self.title_changed.emit(self.overview_id, new_title)
            self.original_title = new_title
            self.overview_title = new_title


class TutorialDialog(QDialog):
    """Interactive tutorial dialog for the dashboard."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Dashboard Tutorial")
        self.setModal(True)
        self.setMinimumSize(500, 400)
        self.current_step = 0
        self.steps = [
            {
                "title": "Welcome to Inclusive Design Wizard!",
                "content": (
                    "This tutorial will guide you through the main features of the dashboard.\n\n"
                    "The Inclusive Design Wizard helps educators create accessible learning "
                    "experiences using UDL (Universal Design for Learning) and WCAG "
                    "(Web Content Accessibility Guidelines) frameworks.\n\n"
                    "Click 'Next' to continue."
                )
            },
            {
                "title": "Starting a New Consultation",
                "content": (
                    "Click the 'Start New Consultation' button to begin a new accessibility review.\n\n"
                    "You'll be asked to:\n"
                    "  1. Give your consultation a name\n"
                    "  2. Select the type of consultation\n\n"
                    "The AI assistant will then guide you through a series of questions "
                    "to help identify accessibility improvements for your learning materials."
                )
            },
            {
                "title": "Recent Consultations",
                "content": (
                    "Your previous consultations appear in the 'Recent Consultations' section.\n\n"
                    "Each card shows:\n"
                    "  - The consultation title (click to edit)\n"
                    "  - The consultation type\n"
                    "  - Current status (In Progress or Complete)\n"
                    "  - Last modified date\n\n"
                    "Click 'Open' to continue a consultation, or 'Delete' to remove it."
                )
            },
            {
                "title": "Settings",
                "content": (
                    "Click the 'Settings' button to configure your AI provider.\n\n"
                    "You can choose between:\n"
                    "  - Local AI (Ollama, LM Studio) - runs on your computer, more private\n"
                    "  - Cloud AI (OpenAI, Anthropic) - requires API key, more powerful\n\n"
                    "Use 'Test Connection' to verify your AI is working before starting."
                )
            },
            {
                "title": "You're Ready!",
                "content": (
                    "That's everything you need to get started!\n\n"
                    "Tips for best results:\n"
                    "  - Be specific about your learning context\n"
                    "  - Describe your learners' needs\n"
                    "  - Ask follow-up questions if needed\n\n"
                    "The AI will provide recommendations based on UDL principles and "
                    "WCAG guidelines to help make your content more accessible.\n\n"
                    "Click 'Finish' to close this tutorial."
                )
            }
        ]
        self.setup_ui()

    def setup_ui(self):
        """Set up the tutorial UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(24, 24, 24, 24)

        # Step indicator
        self.step_indicator = QLabel()
        self.step_indicator.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 12px;")
        layout.addWidget(self.step_indicator)

        # Title
        self.title_label = QLabel()
        self.title_label.setFont(QFont("Arial", 20, QFont.Weight.Bold))
        self.title_label.setWordWrap(True)
        layout.addWidget(self.title_label)

        # Content
        self.content_label = QLabel()
        self.content_label.setFont(QFont("Arial", 14))
        self.content_label.setWordWrap(True)
        self.content_label.setStyleSheet(f"color: {COLORS['text']}; line-height: 1.5;")
        self.content_label.setAlignment(Qt.AlignmentFlag.AlignTop)
        layout.addWidget(self.content_label, 1)

        # Navigation buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        self.prev_btn = QPushButton("Previous")
        self.prev_btn.setProperty("class", "secondary")
        self.prev_btn.clicked.connect(self.prev_step)
        self.prev_btn.setMinimumHeight(44)
        btn_layout.addWidget(self.prev_btn)

        btn_layout.addStretch()

        self.skip_btn = QPushButton("Skip Tutorial")
        self.skip_btn.setProperty("class", "text")
        self.skip_btn.clicked.connect(self.reject)
        self.skip_btn.setMinimumHeight(44)
        btn_layout.addWidget(self.skip_btn)

        self.next_btn = QPushButton("Next")
        self.next_btn.clicked.connect(self.next_step)
        self.next_btn.setMinimumHeight(44)
        self.next_btn.setMinimumWidth(100)
        btn_layout.addWidget(self.next_btn)

        layout.addLayout(btn_layout)

        self.update_content()

    def update_content(self):
        """Update the dialog content for current step."""
        step = self.steps[self.current_step]
        self.step_indicator.setText(f"Step {self.current_step + 1} of {len(self.steps)}")
        self.title_label.setText(step["title"])
        self.content_label.setText(step["content"])

        # Update button states
        self.prev_btn.setEnabled(self.current_step > 0)

        if self.current_step == len(self.steps) - 1:
            self.next_btn.setText("Finish")
        else:
            self.next_btn.setText("Next")

    def next_step(self):
        """Go to next step or finish."""
        if self.current_step < len(self.steps) - 1:
            self.current_step += 1
            self.update_content()
        else:
            self.accept()

    def prev_step(self):
        """Go to previous step."""
        if self.current_step > 0:
            self.current_step -= 1
            self.update_content()


class SessionCard(QFrame):
    """Card widget for displaying a session."""

    open_requested = pyqtSignal(int)
    delete_requested = pyqtSignal(int)
    title_changed = pyqtSignal(int, str)
    notes_requested = pyqtSignal(int, str)  # session_id, session_title

    def __init__(self, session: Session):
        super().__init__()
        self.session_id = session.id
        self.session_title = session.title or "Untitled Consultation"
        self.original_title = self.session_title
        self.setup_ui(session)
        self.setFocusPolicy(Qt.FocusPolicy.TabFocus)

    def setup_ui(self, session: Session):
        """Set up the card UI."""
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {COLORS['dark_card']};
                border-radius: 12px;
                padding: 16px;
            }}
            QFrame:hover {{
                background-color: {COLORS['dark_input']};
            }}
        """)
        self.setMinimumWidth(250)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        layout = QVBoxLayout(self)
        layout.setSpacing(8)
        layout.setContentsMargins(0, 0, 0, 0)

        # Content area - holds all text elements
        content_widget = QWidget()
        content_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(6)
        content_layout.setContentsMargins(0, 0, 0, 0)

        # Editable Title
        self.title_input = QLineEdit(session.title or "Untitled Consultation")
        self.title_input.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        self.title_input.setStyleSheet(f"""
            QLineEdit {{
                background-color: {COLORS['dark_input']};
                color: {COLORS['text']};
                border: 1px solid {COLORS['dark_input']};
                border-radius: 6px;
                padding: 6px 8px;
            }}
            QLineEdit:focus {{
                border: 1px solid {COLORS['primary']};
            }}
        """)
        self.title_input.setAccessibleName("Edit consultation title")
        self.title_input.editingFinished.connect(self.on_title_changed)
        content_layout.addWidget(self.title_input)

        # Type badge
        type_name = CONSULTATION_TYPES.get(session.template_type, {}).get("name", "Custom")
        type_label = QLabel(type_name)
        type_label.setStyleSheet(f"""
            background-color: {COLORS['secondary']};
            color: {COLORS['text']};
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 11px;
        """)
        type_label.setMaximumWidth(120)
        type_label.setSizePolicy(QSizePolicy.Policy.Maximum, QSizePolicy.Policy.Fixed)
        content_layout.addWidget(type_label)

        # Status
        status_text = "Complete" if session.completed else f"In Progress - {session.current_phase}"
        status = QLabel(status_text)
        status.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 12px;")
        status.setWordWrap(True)
        content_layout.addWidget(status)

        # Last modified
        if session.updated_at:
            modified = session.updated_at.strftime("%b %d, %Y at %I:%M %p")
        else:
            modified = "Unknown"
        modified_label = QLabel(f"Last modified: {modified}")
        modified_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px;")
        modified_label.setWordWrap(True)
        content_layout.addWidget(modified_label)

        layout.addWidget(content_widget)

        # Spacer to push buttons to bottom
        layout.addStretch(1)

        # Actions row - separate widget with fixed height
        actions_widget = QWidget()
        actions_widget.setFixedHeight(28)
        actions_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        actions = QHBoxLayout(actions_widget)
        actions.setSpacing(6)
        actions.setContentsMargins(0, 4, 0, 0)

        # Open button (green)
        open_btn = QPushButton("Open")
        open_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {COLORS['success']};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 2px 8px;
                font-size: 11px;
                font-weight: bold;
                min-height: 20px;
                max-height: 24px;
                min-width: 40px;
                max-width: 50px;
            }}
            QPushButton:hover {{
                background-color: #27ae60;
            }}
        """)
        open_btn.setFixedSize(50, 24)
        open_btn.clicked.connect(lambda: self.open_requested.emit(self.session_id))
        open_btn.setAccessibleName(f"Open {session.title}")
        actions.addWidget(open_btn)

        # Notes button
        notes_btn = QPushButton("Notes")
        notes_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {COLORS['primary']};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 2px 8px;
                font-size: 11px;
                font-weight: bold;
                min-height: 20px;
                max-height: 24px;
                min-width: 45px;
                max-width: 55px;
            }}
            QPushButton:hover {{
                background-color: #5849c4;
            }}
        """)
        notes_btn.setFixedSize(55, 24)
        notes_btn.clicked.connect(lambda: self.notes_requested.emit(self.session_id, self.session_title))
        notes_btn.setAccessibleName(f"Notes for {session.title}")
        actions.addWidget(notes_btn)

        actions.addStretch()

        # Delete button (smaller text)
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {COLORS['error']};
                border: 1px solid {COLORS['error']};
                border-radius: 4px;
                padding: 2px 6px;
                font-size: 10px;
                min-height: 18px;
                max-height: 22px;
                min-width: 40px;
                max-width: 50px;
            }}
            QPushButton:hover {{
                background-color: {COLORS['error']};
                color: white;
            }}
        """)
        delete_btn.setFixedSize(50, 22)
        delete_btn.clicked.connect(lambda: self.delete_requested.emit(self.session_id))
        delete_btn.setAccessibleName(f"Delete {session.title}")
        actions.addWidget(delete_btn)

        layout.addWidget(actions_widget)

        # Accessibility
        self.setAccessibleName(f"{session.title}. {type_name}. {status_text}")
        self.setAccessibleDescription("Use Open button to open this consultation")

    def on_title_changed(self):
        """Handle title edit completion."""
        new_title = self.title_input.text().strip()
        if new_title and new_title != self.original_title:
            self.title_changed.emit(self.session_id, new_title)
            self.original_title = new_title


class NewSessionDialog(QDialog):
    """Dialog for creating a new session."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Start New Consultation")
        self.setModal(True)
        self.setMinimumWidth(400)
        self.setup_ui()

    def setup_ui(self):
        """Set up the dialog UI."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {COLORS['dark_card']};
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Title
        title = QLabel("Start New Consultation")
        title.setFont(QFont("Arial", 20, QFont.Weight.Bold))
        layout.addWidget(title)

        # Session name
        name_label = QLabel("Session Name")
        layout.addWidget(name_label)

        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Enter a name for this consultation")
        self.name_input.setAccessibleName("Session name input")
        self.name_input.setMinimumHeight(48)
        layout.addWidget(self.name_input)

        # Consultation type
        type_label = QLabel("Consultation Type")
        layout.addWidget(type_label)

        self.type_combo = QComboBox()
        self.type_combo.setAccessibleName("Select consultation type")
        self.type_combo.setMinimumHeight(48)

        for key, value in CONSULTATION_TYPES.items():
            self.type_combo.addItem(value["name"], key)

        layout.addWidget(self.type_combo)

        # Type description
        self.type_description = QLabel()
        self.type_description.setWordWrap(True)
        self.type_description.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 14px;")
        layout.addWidget(self.type_description)

        self.type_combo.currentIndexChanged.connect(self.update_type_description)
        self.update_type_description()

        layout.addStretch()

        # Buttons
        btn_layout = QHBoxLayout()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setProperty("class", "secondary")
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setMinimumHeight(48)
        btn_layout.addWidget(cancel_btn)

        start_btn = QPushButton("Start Consultation")
        start_btn.clicked.connect(self.accept)
        start_btn.setMinimumHeight(48)
        btn_layout.addWidget(start_btn)

        layout.addLayout(btn_layout)

    def update_type_description(self):
        """Update the description based on selected type."""
        type_key = self.type_combo.currentData()
        if type_key:
            prompt = CONSULTATION_TYPES.get(type_key, {}).get("prompt", "")
            # Get first sentence
            desc = prompt.split(".")[0] + "." if prompt else ""
            self.type_description.setText(desc)

    def get_values(self) -> tuple[str, str]:
        """Get the entered values."""
        name = self.name_input.text().strip() or "New Consultation"
        type_key = self.type_combo.currentData()
        return name, type_key


class DashboardWidget(QWidget):
    """Main dashboard widget."""

    open_session = pyqtSignal(int)
    new_session = pyqtSignal(str, str)
    quick_action = pyqtSignal(str, str)  # For temporary sessions (not auto-saved)
    open_settings = pyqtSignal()
    logout = pyqtSignal()

    def __init__(self, db_manager: DatabaseManager, auth_manager: AuthManager, ai_manager=None):
        super().__init__()
        self.db = db_manager
        self.auth = auth_manager
        self.ai = ai_manager
        self.setup_ui()

    def setup_ui(self):
        """Set up the dashboard UI."""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Main scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setStyleSheet(f"""
            QScrollArea {{
                border: none;
                background-color: transparent;
            }}
            QScrollBar:vertical {{
                background-color: {COLORS['dark_bg']};
                width: 10px;
                border-radius: 5px;
            }}
            QScrollBar::handle:vertical {{
                background-color: {COLORS['dark_input']};
                border-radius: 5px;
                min-height: 20px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: {COLORS['primary']};
            }}
        """)

        scroll_content = QWidget()
        layout = QVBoxLayout(scroll_content)
        layout.setContentsMargins(32, 24, 32, 24)
        layout.setSpacing(24)

        # Header
        header = QHBoxLayout()

        title_section = QVBoxLayout()
        title = QLabel("Dashboard")
        title.setFont(QFont("Arial", 28, QFont.Weight.Bold))
        title.setAccessibleName("Dashboard heading")
        title_section.addWidget(title)

        user = self.auth.get_current_user()
        if user:
            welcome = QLabel(f"Welcome back, {user.email or 'Local User'}")
            welcome.setStyleSheet(f"color: {COLORS['text_muted']};")
            title_section.addWidget(welcome)

        header.addLayout(title_section)
        header.addStretch()

        # Header buttons
        tutorial_btn = QPushButton("Tutorial")
        tutorial_btn.setProperty("class", "secondary")
        tutorial_btn.setAccessibleName("Open tutorial")
        tutorial_btn.setToolTip("Learn how to use the dashboard")
        tutorial_btn.clicked.connect(self.show_tutorial)
        tutorial_btn.setMinimumHeight(44)
        header.addWidget(tutorial_btn)

        overview_btn = QPushButton("Consult Overview")
        overview_btn.setProperty("class", "secondary")
        overview_btn.setAccessibleName("View overview of all consultations")
        overview_btn.setToolTip("AI analysis of your learning progress and patterns")
        overview_btn.clicked.connect(self.show_consult_overview)
        overview_btn.setMinimumHeight(44)
        header.addWidget(overview_btn)

        settings_btn = QPushButton("Settings")
        settings_btn.setProperty("class", "secondary")
        settings_btn.setAccessibleName("Open settings")
        settings_btn.clicked.connect(self.open_settings.emit)
        settings_btn.setMinimumHeight(44)
        header.addWidget(settings_btn)

        logout_btn = QPushButton("Log Out")
        logout_btn.setProperty("class", "text")
        logout_btn.clicked.connect(self.logout.emit)
        logout_btn.setMinimumHeight(44)
        header.addWidget(logout_btn)

        layout.addLayout(header)

        # Quick actions section
        actions_label = QLabel("Quick Actions")
        actions_label.setFont(QFont("Arial", 18, QFont.Weight.Bold))
        layout.addWidget(actions_label)

        actions_layout = QHBoxLayout()
        actions_layout.setSpacing(16)

        quick_actions = [
            ("Start New Consultation", "custom", "Begin a comprehensive accessibility review"),
        ]

        for name, type_key, tooltip in quick_actions:
            btn = QPushButton(name)
            btn.setToolTip(tooltip)
            btn.setAccessibleDescription(tooltip)
            btn.setMinimumHeight(60)
            btn.setMinimumWidth(180)

            if type_key == "custom":
                # Primary button for main action
                pass
            else:
                btn.setProperty("class", "secondary")

            btn.clicked.connect(lambda checked, t=type_key, n=name: self.start_quick_session(t, n))
            actions_layout.addWidget(btn)

        actions_layout.addStretch()
        layout.addLayout(actions_layout)

        # Recent sessions section
        sessions_header = QHBoxLayout()

        sessions_label = QLabel("Recent Consultations")
        sessions_label.setFont(QFont("Arial", 18, QFont.Weight.Bold))
        sessions_header.addWidget(sessions_label)

        layout.addLayout(sessions_header)

        # Sessions scroll area
        self.sessions_area = QScrollArea()
        self.sessions_area.setWidgetResizable(True)
        self.sessions_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.sessions_area.setAccessibleName("Recent consultation sessions")
        self.sessions_area.setFocusPolicy(Qt.FocusPolicy.TabFocus)

        self.sessions_container = QWidget()
        self.sessions_layout = QGridLayout(self.sessions_container)
        self.sessions_layout.setSpacing(16)
        # Set columns to stretch equally
        self.sessions_layout.setColumnStretch(0, 1)
        self.sessions_layout.setColumnStretch(1, 1)
        self.sessions_layout.setColumnStretch(2, 1)

        self.sessions_area.setWidget(self.sessions_container)
        layout.addWidget(self.sessions_area)

        # Personal Consult Overview section
        layout.addSpacing(24)

        overviews_header = QHBoxLayout()

        overviews_label = QLabel("Personal Consult Overviews")
        overviews_label.setFont(QFont("Arial", 18, QFont.Weight.Bold))
        overviews_header.addWidget(overviews_label)

        layout.addLayout(overviews_header)

        # Overviews scroll area
        self.overviews_area = QScrollArea()
        self.overviews_area.setWidgetResizable(True)
        self.overviews_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.overviews_area.setAccessibleName("Personal consultation overviews")
        self.overviews_area.setFocusPolicy(Qt.FocusPolicy.TabFocus)

        self.overviews_container = QWidget()
        self.overviews_layout = QGridLayout(self.overviews_container)
        self.overviews_layout.setSpacing(16)
        self.overviews_layout.setColumnStretch(0, 1)
        self.overviews_layout.setColumnStretch(1, 1)
        self.overviews_layout.setColumnStretch(2, 1)

        self.overviews_area.setWidget(self.overviews_container)
        layout.addWidget(self.overviews_area)

        # Add scroll content to main scroll area
        scroll.setWidget(scroll_content)
        main_layout.addWidget(scroll)

        self.load_sessions()
        self.load_overviews()

    def load_sessions(self):
        """Load and display user sessions."""
        # Clear existing cards
        while self.sessions_layout.count():
            item = self.sessions_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        user = self.auth.get_current_user()
        if not user:
            return

        sessions = self.db.get_user_sessions(user.id)

        if not sessions:
            empty_label = QLabel("No consultations yet. Start one above!")
            empty_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 16px;")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.sessions_layout.addWidget(empty_label, 0, 0)
            return

        # Display in grid (3 columns)
        for i, session in enumerate(sessions[:12]):  # Show last 12
            row = i // 3
            col = i % 3

            card = SessionCard(session)
            card.open_requested.connect(self.open_session.emit)
            card.delete_requested.connect(self.confirm_delete_session)
            card.title_changed.connect(self.save_session_title)
            card.notes_requested.connect(self.show_notes_dialog)

            self.sessions_layout.addWidget(card, row, col)

    def start_quick_session(self, type_key: str, name: str):
        """Start a quick session with preset type (temporary, not auto-saved)."""
        if type_key == "custom":
            self.show_new_session_dialog()
        else:
            title = f"{name} - {datetime.now().strftime('%b %d, %Y')}"
            self.quick_action.emit(title, type_key)

    def show_new_session_dialog(self):
        """Show dialog to create new session."""
        dialog = NewSessionDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            title, type_key = dialog.get_values()
            self.new_session.emit(title, type_key)

    def confirm_delete_session(self, session_id: int):
        """Confirm and delete a session."""
        reply = QMessageBox.question(
            self,
            "Delete Consultation",
            "Are you sure you want to delete this consultation? This action cannot be undone.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.db.delete_session(session_id)
            self.load_sessions()

    def save_session_title(self, session_id: int, new_title: str):
        """Save updated session title."""
        self.db.update_session(session_id, title=new_title)

    def refresh(self):
        """Refresh the dashboard."""
        self.load_sessions()
        self.load_overviews()

    def show_tutorial(self):
        """Show the dashboard tutorial."""
        dialog = TutorialDialog(self)
        dialog.exec()

    def show_notes_dialog(self, session_id: int, session_title: str):
        """Show the notes dialog for a session."""
        notes = self.db.get_session_notes(session_id)
        dialog = NotesDialog(session_id, session_title, notes, self)
        dialog.note_added.connect(self.save_note)
        dialog.exec()

    def save_note(self, session_id: int, note_content: str):
        """Save a note to a session."""
        self.db.add_session_note(session_id, note_content)
        QMessageBox.information(
            self,
            "Note Saved",
            "Your note has been saved with a timestamp."
        )

    def show_consult_overview(self):
        """Show the AI-generated consultation overview."""
        if not self.ai:
            QMessageBox.warning(
                self,
                "AI Not Configured",
                "Please configure your AI settings first."
            )
            return

        dialog = ConsultOverviewDialog(self.db, self.auth, self.ai, self)
        dialog.exec()
        # Refresh overviews list after dialog closes
        self.load_overviews()

    def load_overviews(self):
        """Load and display user overviews."""
        # Clear existing cards
        while self.overviews_layout.count():
            item = self.overviews_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        user = self.auth.get_current_user()
        if not user:
            return

        overviews = self.db.get_user_overviews(user.id)

        if not overviews:
            empty_label = QLabel("No overviews yet. Click 'Consult Overview' to generate one!")
            empty_label.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 14px;")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.overviews_layout.addWidget(empty_label, 0, 0)
            return

        # Display in grid (3 columns)
        for i, overview in enumerate(overviews[:9]):  # Show last 9
            row = i // 3
            col = i % 3

            card = OverviewCard(overview)
            card.open_requested.connect(self.open_overview)
            card.delete_requested.connect(self.confirm_delete_overview)
            card.notes_requested.connect(self.show_overview_notes_dialog)
            card.title_changed.connect(self.save_overview_title)

            self.overviews_layout.addWidget(card, row, col)

    def open_overview(self, overview_id: int):
        """Open a saved overview."""
        overview = self.db.get_overview_by_id(overview_id)
        if overview:
            dialog = ViewOverviewDialog(overview, self)
            dialog.exec()

    def confirm_delete_overview(self, overview_id: int):
        """Confirm and delete an overview."""
        reply = QMessageBox.question(
            self,
            "Delete Overview",
            "Are you sure you want to delete this overview? This action cannot be undone.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.db.delete_overview(overview_id)
            self.load_overviews()

    def show_overview_notes_dialog(self, overview_id: int, overview_title: str):
        """Show the notes dialog for an overview."""
        notes = self.db.get_overview_notes(overview_id)
        dialog = OverviewNotesDialog(overview_id, overview_title, notes, self)
        dialog.note_added.connect(self.save_overview_note)
        dialog.exec()

    def save_overview_note(self, overview_id: int, note_content: str):
        """Save a note to an overview."""
        self.db.add_overview_note(overview_id, note_content)
        QMessageBox.information(
            self,
            "Note Saved",
            "Your note has been saved with a timestamp."
        )

    def save_overview_title(self, overview_id: int, new_title: str):
        """Save updated overview title."""
        self.db.update_overview_title(overview_id, new_title)
