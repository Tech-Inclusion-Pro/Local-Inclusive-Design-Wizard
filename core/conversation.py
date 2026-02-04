"""Conversation and session management."""

from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import PHASES, PHASE_ORDER
from prompts.system_prompts import get_phase_reasoning, CONSULTATION_TYPES
from .database import DatabaseManager, Session


class ConversationManager:
    """Manage consultation conversations and progress."""

    def __init__(self, db_manager: DatabaseManager):
        self.db = db_manager
        self.current_session: Optional[Session] = None
        self.current_phase_index = 0
        self.current_question_index = 0
        self.is_temporary = False
        self._temp_user_id = None
        self._temp_conversation = []

    def start_new_session(self, user_id: int, title: str, template_type: str) -> Session:
        """Start a new consultation session."""
        session = self.db.create_session(user_id, title, template_type)
        self.current_session = session
        self.current_phase_index = 0
        self.current_question_index = 0
        self.is_temporary = False
        return session

    def start_temporary_session(self, user_id: int, title: str, template_type: str):
        """Start a temporary session without saving to database."""
        # Create a temporary session object (not saved to DB)
        self.current_session = Session(
            user_id=user_id,
            title=title,
            template_type=template_type,
            current_phase=PHASE_ORDER[0] if PHASE_ORDER else "discovery",
            current_question_index=0
        )
        self.current_phase_index = 0
        self.current_question_index = 0
        self.is_temporary = True
        self._temp_user_id = user_id
        self._temp_conversation = []

    def save_temporary_session(self) -> Optional[Session]:
        """Save a temporary session to the database."""
        if not self.is_temporary or not self.current_session:
            return self.current_session

        # Create the session in database
        session = self.db.create_session(
            self._temp_user_id,
            self.current_session.title,
            self.current_session.template_type
        )

        # Copy conversation messages
        db_session = self.db.get_session()
        try:
            db_sess = db_session.query(Session).filter(Session.id == session.id).first()
            if db_sess:
                for msg in self._temp_conversation:
                    db_sess.add_message(msg["role"], msg["content"], msg.get("reasoning"))
                db_sess.current_phase = PHASE_ORDER[self.current_phase_index] if self.current_phase_index < len(PHASE_ORDER) else "complete"
                db_sess.current_question_index = self.current_question_index
                db_session.commit()
                self.current_session = db_sess
        finally:
            db_session.close()

        self.is_temporary = False
        self._temp_conversation = []
        return self.current_session

    def load_session(self, session_id: int) -> Optional[Session]:
        """Load an existing session."""
        session = self.db.get_session_by_id(session_id)
        if session:
            self.current_session = session
            # Restore position
            if session.current_phase in PHASE_ORDER:
                self.current_phase_index = PHASE_ORDER.index(session.current_phase)
            self.current_question_index = session.current_question_index
            self.is_temporary = False
            self._temp_conversation = []
        return session

    def get_current_phase(self) -> dict:
        """Get current phase information."""
        if self.current_phase_index >= len(PHASE_ORDER):
            return {
                "key": "complete",
                "name": "Consultation Complete",
                "questions": []
            }

        phase_key = PHASE_ORDER[self.current_phase_index]
        phase_data = PHASES[phase_key]
        return {
            "key": phase_key,
            "name": phase_data["name"],
            "questions": phase_data["questions"]
        }

    def get_current_question(self) -> Optional[str]:
        """Get the current question to ask."""
        phase = self.get_current_phase()
        if self.current_question_index < len(phase["questions"]):
            return phase["questions"][self.current_question_index]
        return None

    def get_phase_reasoning(self) -> dict:
        """Get reasoning for current phase."""
        if self.current_phase_index >= len(PHASE_ORDER):
            return {}
        phase_key = PHASE_ORDER[self.current_phase_index]
        return get_phase_reasoning(phase_key, self.current_question_index)

    def advance_question(self) -> bool:
        """Move to next question. Returns True if moved to new phase."""
        if not self.current_session:
            return False

        phase = self.get_current_phase()
        self.current_question_index += 1

        new_phase = False
        if self.current_question_index >= len(phase["questions"]):
            # Move to next phase
            self.current_phase_index += 1
            self.current_question_index = 0
            new_phase = True

            if self.current_phase_index >= len(PHASE_ORDER):
                # Consultation complete - only update DB if not temporary
                if not getattr(self, 'is_temporary', False):
                    self.db.update_session(
                        self.current_session.id,
                        completed=True,
                        current_phase="complete",
                        current_question_index=0
                    )
                return new_phase

        # Update session in DB only if not temporary
        if not getattr(self, 'is_temporary', False):
            current_phase_key = PHASE_ORDER[self.current_phase_index] if self.current_phase_index < len(PHASE_ORDER) else "complete"
            self.db.update_session(
                self.current_session.id,
                current_phase=current_phase_key,
                current_question_index=self.current_question_index
            )

        return new_phase

    def add_message(self, role: str, content: str, reasoning: dict = None):
        """Add message to current session."""
        if not self.current_session:
            return

        # Handle temporary sessions
        if getattr(self, 'is_temporary', False):
            self._temp_conversation.append({
                "role": role,
                "content": content,
                "reasoning": reasoning,
                "timestamp": datetime.now().isoformat()
            })
            return

        # Reload session to get latest conversation
        db_session = self.db.get_session()
        try:
            session = db_session.query(Session).filter(
                Session.id == self.current_session.id
            ).first()
            if session:
                session.add_message(role, content, reasoning)
                db_session.commit()
                self.current_session = session
        finally:
            db_session.close()

    def get_conversation(self) -> list:
        """Get all messages in current session."""
        if not self.current_session:
            return []

        # Handle temporary sessions
        if getattr(self, 'is_temporary', False):
            return self._temp_conversation

        # Reload to get latest
        session = self.db.get_session_by_id(self.current_session.id)
        return session.conversation if session else []

    def get_progress(self) -> dict:
        """Get consultation progress information."""
        total_questions = sum(len(p["questions"]) for p in PHASES.values())
        completed_questions = 0

        for i, phase_key in enumerate(PHASE_ORDER):
            if i < self.current_phase_index:
                completed_questions += len(PHASES[phase_key]["questions"])
            elif i == self.current_phase_index:
                completed_questions += self.current_question_index

        progress_percent = (completed_questions / total_questions * 100) if total_questions > 0 else 0

        phases_status = []
        for i, phase_key in enumerate(PHASE_ORDER):
            phase = PHASES[phase_key]
            if i < self.current_phase_index:
                status = "completed"
            elif i == self.current_phase_index:
                status = "current"
            else:
                status = "pending"

            phases_status.append({
                "key": phase_key,
                "name": phase["name"],
                "status": status
            })

        return {
            "percent": progress_percent,
            "completed_questions": completed_questions,
            "total_questions": total_questions,
            "current_phase": self.current_phase_index,
            "current_question": self.current_question_index,
            "phases": phases_status
        }

    def is_complete(self) -> bool:
        """Check if consultation is complete."""
        return self.current_phase_index >= len(PHASE_ORDER)

    def export_to_docx(self, filepath: str) -> bool:
        """Export consultation to accessible DOCX format."""
        if not self.current_session:
            return False

        try:
            doc = Document()

            # Set document properties
            doc.core_properties.title = self.current_session.title
            doc.core_properties.subject = "Inclusive Design Consultation"

            # Title
            title = doc.add_heading(self.current_session.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Consultation Type: {CONSULTATION_TYPES.get(self.current_session.template_type, {}).get('name', 'Custom')}\n")
            meta.add_run(f"Created: {self.current_session.created_at.strftime('%Y-%m-%d')}\n")
            meta.add_run(f"Status: {'Complete' if self.current_session.completed else 'In Progress'}")

            doc.add_paragraph()

            # Summary section
            doc.add_heading("Summary", level=1)
            doc.add_paragraph("This document contains a record of the inclusive design consultation, organized by UDL and WCAG frameworks.")

            # Conversation by phase
            conversation = self.get_conversation()
            current_phase = None

            for msg in conversation:
                # Add phase headers as we go through conversation
                phase_info = msg.get("reasoning", {}).get("phase")
                if phase_info and phase_info != current_phase:
                    current_phase = phase_info
                    phase_name = PHASES.get(phase_info, {}).get("name", phase_info)
                    doc.add_heading(phase_name, level=1)

                # Add message
                role_label = "Consultant" if msg["role"] == "assistant" else "Educator"
                p = doc.add_paragraph()
                role_run = p.add_run(f"{role_label}: ")
                role_run.bold = True
                p.add_run(msg["content"])

                # Add reasoning if present
                reasoning = msg.get("reasoning", {})
                if reasoning and msg["role"] == "assistant":
                    if reasoning.get("framework"):
                        r_para = doc.add_paragraph()
                        r_para.style = "Quote"
                        r_para.add_run(f"Framework: {reasoning['framework']}")
                        if reasoning.get("why"):
                            r_para.add_run(f"\nWhy: {reasoning['why']}")

            # Recommendations section
            doc.add_heading("Key Recommendations", level=1)
            doc.add_paragraph("Based on this consultation, consider the following action items:")

            # Extract any recommendations from assistant messages
            for msg in conversation:
                if msg["role"] == "assistant" and "recommend" in msg["content"].lower():
                    p = doc.add_paragraph(style="List Bullet")
                    # Take first sentence containing recommendation
                    sentences = msg["content"].split(".")
                    for s in sentences:
                        if "recommend" in s.lower():
                            p.add_run(s.strip() + ".")
                            break

            # Accessibility statement
            doc.add_paragraph()
            doc.add_heading("Accessibility", level=1)
            doc.add_paragraph(
                "This document was generated with accessibility in mind. "
                "It uses heading styles for navigation, semantic structure, "
                "and avoids images without alternative text."
            )

            # Save
            doc.save(filepath)
            return True

        except Exception as e:
            print(f"Export error: {e}")
            return False
